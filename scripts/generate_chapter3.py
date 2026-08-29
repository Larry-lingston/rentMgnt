"""Generate Chapter 3 docx from Rent Manager README / implementation."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"c:\Users\Larry Lingston\Desktop\Chapter_Three_Design_and_Methodology.docx"

doc = Document()

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

# Title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("CHAPTER THREE")
r.bold = True
r.font.size = Pt(14)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("DESIGN AND METHODOLOGY")
r2.bold = True
r2.font.size = Pt(14)

add_heading("3.1 Introduction", level=2)
add_para(
    "This chapter describes the design and methodology used to develop the Rent Manager "
    "mobile rent management system. It outlines the proposed system, functional and "
    "non-functional requirements, technologies selected, system architecture, database "
    "design, and modelling approach. The system was built as a full-stack mobile "
    "application that connects landlords, tenants, maintenance crew, and room seekers "
    "through a single platform, with all monetary values recorded in Ghana Cedis (GHS)."
)

add_heading("3.2 The Proposed System", level=2)
add_para(
    "Rent Manager is a mobile-based rent management application designed to replace "
    "manual record keeping with a centralized digital system. The solution consists of "
    "two main parts: an Expo mobile client and a Node.js REST API backed by SQLite."
)
add_para(
    "The system supports four user roles. Landlords (admin) manage properties, tenants, "
    "payments, booking requests, maintenance, crew members, and reports. Tenants view "
    "rent balances, pay rent in-app (simulated), submit maintenance requests, and receive "
    "alerts. Maintenance crew claim open jobs or update assigned tasks. Room seekers browse "
    "vacant listings without logging in, request to book a room, and become tenants after "
    "landlord approval."
)
add_para(
    "Key capabilities include property and room management with photos (camera or gallery), "
    "property map locations, rent notifications, digital receipts, booking workflow, and "
    "role-based dashboards. Compared with paper-based methods, the system improves data "
    "accuracy, access speed, and communication between stakeholders."
)

add_heading("3.3 System Requirements", level=2)

add_heading("3.3.1 Functional Requirements", level=3)
add_para("The system shall provide the following services:")

add_para("Authentication and user management")
add_bullets([
    "Register accounts as landlord, tenant, maintenance crew, or room seeker.",
    "Log in with username and password; route users to the correct portal by role.",
    "Update profile details and change password.",
    "Link tenant and maintenance accounts to a landlord using the landlord username.",
])

add_para("Property and room management")
add_bullets([
    "Create, view, update, and delete properties with address, type, description, and optional map coordinates.",
    "Upload property photos from the device camera or gallery.",
    "Track rooms per property with rent amount, occupancy status (vacant/occupied), and images.",
    "Display properties on a map (react-native-maps on mobile).",
])

add_para("Tenant and booking management")
add_bullets([
    "Landlords add and manage tenants; assign tenants to vacant rooms.",
    "Public browse screen lists vacant rooms with photos and rent in GHS.",
    "Seekers submit booking requests; landlords approve or reject them.",
    "Approved bookings create a tenant profile and assign the room automatically.",
])

add_para("Rent payment management")
add_bullets([
    "Landlords record payments (cash, bank transfer, check, mobile money) and view outstanding balances.",
    "Tenants pay rent in-app (simulated checkout) and view payment history and receipts.",
    "Generate monthly collection, outstanding balance, occupancy, and tenant reports.",
])

add_para("Maintenance management")
add_bullets([
    "Tenants and landlords submit maintenance requests with title and description.",
    "Assign jobs as first-come-first-serve (open pool) or to a selected crew member.",
    "Crew members claim open jobs and update task status (pending, completed, cancelled).",
])

add_para("Notifications")
add_bullets([
    "Send rent upcoming and rent due alerts to landlords and tenants.",
    "Display in-app notification lists with read/unread status.",
])

add_heading("3.3.2 Non-Functional Requirements", level=3)
add_bullets([
    "Performance: API responses and dashboard loads should complete within a few seconds on a local network.",
    "Security: Passwords hashed with bcrypt; access protected with JWT; role-based route restrictions on the API.",
    "Usability: Mobile UI with tab navigation, clear labels, and Ghana Cedis (₵) for all amounts.",
    "Maintainability: Separate backend and mobile projects; modular Express routes and Prisma schema.",
    "Portability: Runs on Node.js 18+; mobile client tested via Expo Go on Android/iOS.",
    "Availability: Designed for development deployment; production would require hosted API and database.",
])

add_heading("3.4 Selection of Technologies and Tools", level=2)
add_para(
    "Technologies were chosen for cross-platform mobile delivery, rapid development, and "
    "a lightweight backend suitable for academic and prototype deployment."
)

add_para("Mobile client (app/)")
add_bullets([
    "Expo SDK 53 — development platform and tooling.",
    "React Native 0.79 and React 19 — user interface.",
    "Expo Router — file-based navigation and role-specific tab layouts.",
    "react-native-maps — property location pins.",
    "expo-image-picker — camera and gallery for property photos.",
    "@react-native-async-storage/async-storage — persist JWT on device.",
    "@expo/vector-icons (Ionicons) — icons.",
])

add_para("Backend API (backend/)")
add_bullets([
    "Node.js 18+ — server runtime.",
    "Express 5 — REST API framework.",
    "Prisma 6 — ORM and database migrations.",
    "SQLite — embedded database (dev.db); no separate database server required.",
    "jsonwebtoken — JWT authentication.",
    "bcryptjs — password hashing.",
    "multer — image upload to server storage.",
    "cors — cross-origin access for mobile client.",
])

add_para("Development tools")
add_bullets([
    "Visual Studio Code / Cursor — code editing.",
    "npm — package management.",
    "Expo Go — mobile testing on physical devices.",
    "Prisma Studio (optional) — database inspection.",
])

add_heading("3.5 System Architecture", level=2)
add_para(
    "The system uses a three-tier architecture: presentation (mobile app), application "
    "(Express API), and data (SQLite via Prisma)."
)
add_bullets([
    "Presentation layer: Expo mobile app with screens for login, registration, browse, "
    "and role-based portals (landlord, tenant, maintenance, seeker). Communicates with "
    "the API over HTTP/JSON.",
    "Application layer: Express server on port 3000. Handles authentication, business "
    "logic, file uploads, rent notification generation, and role checks before database access.",
    "Data layer: SQLite file (prisma/dev.db) accessed through Prisma Client. Stores users, "
    "properties, rooms, tenants, payments, bookings, maintenance requests, and notifications.",
])
add_para(
    "Authentication flow: the client sends credentials to /api/auth/login; the server "
    "returns a JWT; subsequent requests include Authorization: Bearer <token>. The mobile "
    "app stores the token in AsyncStorage."
)

add_heading("3.6 System Design", level=2)
add_para(
    "UML-style models were used to describe user interaction and system behaviour. "
    "Diagram placeholders are retained for figures to be inserted in the final document."
)

add_heading("3.6.1 System Flowchart", level=3)
add_para(
    "The user opens the mobile app (splash screen), then logs in or registers. "
    "Based on role, the app routes to: landlord dashboard, tenant home, maintenance "
    "open jobs, or seeker browse. Seekers may browse listings without login. "
    "After landlord approval of a booking, the seeker logs out and back in to access "
    "the tenant portal."
)
add_para("Figure 3.1: System Flowchart")
add_para("(Insert System Flowchart here.)")

add_heading("3.6.2 Use Case Diagram", level=3)
add_para("Major actors and use cases:")
add_bullets([
    "Landlord: manage properties, tenants, payments, bookings, maintenance, crew, reports.",
    "Tenant: view rent summary, pay rent, view receipts, submit maintenance, view alerts.",
    "Maintenance crew: claim open jobs, update task status, view property on map.",
    "Room seeker: browse vacant rooms, request booking, track request status.",
    "Public user: browse available rooms without authentication.",
])
add_para("Figure 3.2: Use Case Diagram")
add_para("(Insert Use Case Diagram here.)")

add_heading("3.6.3 Activity Diagram — Rent Payment", level=3)
add_para(
    "Tenant opens Payments tab, views outstanding balance, taps Pay Rent, enters amount, "
    "and confirms simulated payment. The API records the payment, generates a receipt "
    "number, updates balances, and may notify the landlord."
)
add_para("Figure 3.3: Activity Diagram")
add_para("(Insert Activity Diagram here.)")

add_heading("3.6.4 Sequence Diagram — Rent Payment", level=3)
add_para(
    "Tenant app sends POST /api/tenant-portal/payments/pay with JWT and amount. "
    "Express validates token and role, checks outstanding balance, creates Payment record "
    "in SQLite via Prisma, returns receipt details to the app."
)
add_para("Figure 3.4: Sequence Diagram")
add_para("(Insert Sequence Diagram here.)")

add_heading("3.6.5 State Transition Diagram — Booking Request", level=3)
add_para(
    "A room booking starts as pending when a seeker submits a request. The landlord "
    "approves (status becomes approved; tenant and room assigned) or rejects (status rejected). "
    "Maintenance tasks transition among pending, completed, and cancelled."
)
add_para("Figure 3.5: State Transition Diagram")
add_para("(Insert State Transition Diagram here.)")

add_heading("3.7 Database Design", level=2)
add_para(
    "The database follows a relational model implemented in SQLite and defined in "
    "backend/prisma/schema.prisma. Prisma maps tables to models and enforces "
    "relationships through foreign keys."
)

entities = [
    ("3.7.1 User", "Stores system accounts. Attributes: id, username, email, password (hashed), name, phone, role (admin, tenant, maintenance, seeker), landlordId, tenantProfileId, timestamps. Role determines API and UI access."),
    ("3.7.2 Property", "Stores rental buildings. Attributes: id, name, address, latitude, longitude, type, totalRooms, description, imageUrl, images (JSON array), userId (landlord). One landlord has many properties."),
    ("3.7.3 Room", "Stores units within a property. Attributes: id, roomNumber, rentAmount, status (vacant/occupied), imageUrl, propertyId. One room may have one tenant; supports room bookings."),
    ("3.7.4 Tenant", "Stores tenant profiles linked to a landlord. Attributes: id, name, email, phone, leaseStart, leaseEnd, roomId, userId. Linked to User for tenant login and to Payment records."),
    ("3.7.5 RoomBooking", "Stores seeker booking requests. Attributes: id, roomId, applicantUserId, landlordId, status (pending/approved/rejected), message, timestamps."),
    ("3.7.6 Payment", "Stores rent payments. Attributes: id, tenantId, amount, paymentDate, method, status, transactionRef, receiptNumber, notes."),
    ("3.7.7 MaintenanceRequest", "Stores repair requests. Attributes: id, title, description, status, priority, propertyId, tenantId, assignmentMode (open/selected), assignedToId, requestedBy, timestamps."),
    ("3.7.8 Notification", "Stores in-app alerts. Attributes: id, userId, type, title, message, read, createdAt. Types include rent_due, rent_upcoming, maintenance, and booking."),
]

for title, body in entities:
    add_heading(title, level=3)
    add_para(body)

add_heading("3.8 Entity Relationship Diagram (ERD)", level=2)
add_para("Main relationships:")
add_bullets([
    "One User (landlord) owns many Properties; each Property has many Rooms.",
    "One Room may have one Tenant; one Tenant has many Payments.",
    "One User (seeker) may create many RoomBookings for Rooms owned by a landlord.",
    "MaintenanceRequest links Property, optional Tenant, and optional assigned crew User.",
    "Notification belongs to one User.",
])
add_para("Figure 3.6: Entity Relationship Diagram (ERD)")
add_para("(Insert ER Diagram here.)")

add_heading("3.9 Database Normalization", level=2)
add_para(
    "Tables were normalized to Third Normal Form (3NF). Each entity has a single primary "
    "key (id). Repeating groups (e.g. multiple property images) are stored as a JSON "
    "array in the images field while maintaining one row per property. Foreign keys "
    "link tenants to rooms, rooms to properties, and payments to tenants, reducing "
    "redundancy and supporting consistent updates."
)

add_heading("3.10 Security Considerations", level=2)
add_bullets([
    "Authentication: JWT issued on login/register; required on protected API routes.",
    "Authorization: requireRole middleware restricts endpoints by role (e.g. admin-only property CRUD).",
    "Password storage: bcrypt hashing before saving to the database.",
    "Input validation: required fields checked on registration, property create, and payment endpoints.",
    "File uploads: restricted to image types and size limit (8 MB) via multer.",
    "Development note: JWT_SECRET and HTTPS should be configured for production deployment.",
])

add_heading("3.11 Chapter Summary", level=2)
add_para(
    "This chapter presented the design and methodology for Rent Manager—a mobile rent "
    "management system built with Expo, React Native, Node.js, Express, Prisma, and SQLite. "
    "It defined functional and non-functional requirements, described the three-tier "
    "architecture, outlined database entities aligned with the implementation, and "
    "summarized security measures. The design supports landlords, tenants, maintenance "
    "crew, and room seekers, with rent and reports expressed in Ghana Cedis, providing "
    "a clear blueprint for the implemented application."
)

doc.save(OUT)
print(f"Saved: {OUT}")
